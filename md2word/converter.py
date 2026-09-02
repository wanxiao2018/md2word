"""Markdown to Word converter utilities."""

from __future__ import annotations

import os
import re
import shutil
import subprocess
import tempfile
from dataclasses import dataclass
from pathlib import Path
from typing import Optional

from .docxstyle import polish_docx
from .i18n import t
from .mathprep import plain_text_from_markdown, prepare_markdown_math, strip_thematic_breaks


SUPPORTED_EXTS = {".md", ".markdown", ".txt", ".mdx"}

# Pandoc reader with math-friendly extensions.
# yaml_metadata_block is disabled so leading --- is not eaten as YAML.
_PANDOC_FROM = (
    "markdown"
    "+pipe_tables"
    "+grid_tables"
    "+strikeout"
    "+task_lists"
    "+fenced_code_attributes"
    "+auto_identifiers"
    "+tex_math_dollars"
    "+tex_math_single_backslash"
    "+raw_html"
    "-yaml_metadata_block"
)


@dataclass
class ConvertResult:
    success: bool
    message: str
    output_path: Optional[Path] = None
    content: Optional[str] = None
    engine: str = "pandoc"


_pandoc_cache: Optional[str] = None
_pandoc_checked = False


def find_pandoc() -> Optional[str]:
    """Locate pandoc executable."""
    global _pandoc_cache, _pandoc_checked
    if _pandoc_checked:
        return _pandoc_cache
    found = shutil.which("pandoc")
    if found:
        _pandoc_checked = True
        _pandoc_cache = found
        return found
    candidates = [
        Path(os.environ.get("LOCALAPPDATA", "")) / "Pandoc" / "pandoc.exe",
        Path(r"C:\Program Files\Pandoc\pandoc.exe"),
        Path(r"C:\Program Files (x86)\Pandoc\pandoc.exe"),
        Path("/opt/homebrew/bin/pandoc"),
        Path("/usr/local/bin/pandoc"),
        Path("/opt/local/bin/pandoc"),
    ]
    for path in candidates:
        if path.is_file():
            _pandoc_checked = True
            _pandoc_cache = str(path)
            return _pandoc_cache
    _pandoc_checked = True
    _pandoc_cache = None
    return None


def looks_like_markdown(text: str) -> bool:
    """Heuristic: does this text look like Markdown?"""
    if not text or not text.strip():
        return False
    sample = text[:8000]
    patterns = [
        r"(?m)^#{1,6}\s+\S",  # headings
        r"(?m)^(\s*[-*+]\s+|\s*\d+\.\s+)",  # lists
        r"```",  # fenced code
        r"\*\*[^*\n]+\*\*",  # bold
        r"(?<!\*)\*[^*\n]+\*(?!\*)",  # italic
        r"(?m)^>\s+\S",  # blockquote
        r"\[[^\]]+\]\([^)]+\)",  # links
        r"(?m)^\|.+\|",  # tables
        r"`[^`\n]+`",  # inline code
        r"(?m)^---+\s*$",  # hr
        r"\\\(|\\\[|\$\$|\\alpha|\\beta|\\gamma",  # math-ish
    ]
    hits = sum(1 for p in patterns if re.search(p, sample))
    return hits >= 1


def normalize_markdown(text: str) -> str:
    """Cleanup AI-exported Markdown and wrap math for Word OMML."""
    if not text:
        return ""
    text = text.replace("\r\n", "\n").replace("\r", "\n")
    text = text.replace("\ufeff", "").replace("\u200b", "").replace("\u200c", "")
    # Fix common Chinese-AI typos near math: 攰 → 改 (garbled “改为”)
    text = text.replace("攰为", "改为").replace("攰", "改")
    text = re.sub(r"(?m)^[ \t]+```", "```", text)
    text = re.sub(r"([^\n])\n```", r"\1\n\n```", text)
    # Ensure blank line before ATX headings when glued to previous text
    text = re.sub(r"([^\n])\n(#{1,6}\s+)", r"\1\n\n\2", text)
    text = strip_thematic_breaks(text)
    text = prepare_markdown_math(text)
    return text.strip() + "\n"


def _run_pandoc(
    pandoc: str,
    input_path: Path,
    output_path: Path,
    to_format: str,
    extra_args: Optional[list[str]] = None,
    standalone: bool = True,
) -> None:
    cmd = [
        pandoc,
        str(input_path),
        "-f",
        _PANDOC_FROM,
        "-t",
        to_format,
        "-o",
        str(output_path),
        "--wrap=none",
    ]
    if standalone:
        cmd.append("--standalone")
    if extra_args:
        cmd.extend(extra_args)
    creationflags = 0
    if os.name == "nt":
        creationflags = getattr(subprocess, "CREATE_NO_WINDOW", 0)
    proc = subprocess.run(
        cmd,
        capture_output=True,
        text=True,
        encoding="utf-8",
        errors="replace",
        creationflags=creationflags,
        timeout=60,
    )
    if proc.returncode != 0:
        err = (proc.stderr or proc.stdout or "unknown error").strip()
        raise RuntimeError(f"Pandoc failed: {err}")


def markdown_to_docx(md_text: str, output_path: Path) -> ConvertResult:
    """Convert Markdown text to a styled .docx with Word OMML equations."""
    md_text = normalize_markdown(md_text)
    if not md_text.strip():
        return ConvertResult(False, t("empty"))

    pandoc = find_pandoc()
    output_path = Path(output_path)
    output_path.parent.mkdir(parents=True, exist_ok=True)

    with tempfile.TemporaryDirectory(prefix="md2word_") as tmp:
        tmp_dir = Path(tmp)
        md_path = tmp_dir / "input.md"
        md_path.write_text(md_text, encoding="utf-8")

        if pandoc:
            try:
                _run_pandoc(pandoc, md_path, output_path, "docx")
                polish_docx(output_path)
                return ConvertResult(
                    True,
                    t("saved_path", path=output_path),
                    output_path=output_path,
                    engine="pandoc",
                )
            except Exception as exc:  # noqa: BLE001
                fallback_msg = str(exc)
        else:
            fallback_msg = t("no_pandoc_builtin")

        try:
            _fallback_markdown_to_docx(md_text, output_path)
            polish_docx(output_path)
            return ConvertResult(
                True,
                t("saved_fallback", path=output_path, detail=fallback_msg),
                output_path=output_path,
                engine="python-docx",
            )
        except Exception as exc:  # noqa: BLE001
            return ConvertResult(False, t("convert_fail_exc", exc=exc))


def markdown_to_html(md_text: str) -> ConvertResult:
    """Convert Markdown to HTML string for clipboard fallback / preview."""
    md_text = normalize_markdown(md_text)
    if not md_text.strip():
        return ConvertResult(False, t("empty"))

    pandoc = find_pandoc()
    with tempfile.TemporaryDirectory(prefix="md2word_") as tmp:
        tmp_dir = Path(tmp)
        md_path = tmp_dir / "input.md"
        html_path = tmp_dir / "out.html"
        md_path.write_text(md_text, encoding="utf-8")

        if pandoc:
            try:
                _run_pandoc(
                    pandoc,
                    md_path,
                    html_path,
                    "html5",
                    extra_args=["--mathml"],
                    standalone=False,
                )
                body = html_path.read_text(encoding="utf-8").strip()
                html = (
                    "<!DOCTYPE html><html xmlns:m="
                    "'http://schemas.openxmlformats.org/officeDocument/2006/math'"
                    " xmlns:w='urn:schemas-microsoft-com:office:word'><head>"
                    "<meta http-equiv='Content-Type' content='text/html; charset=utf-8'>"
                    "</head><body>"
                    f"{body}"
                    "</body></html>"
                )
                html = _inject_word_friendly_styles(html)
                return ConvertResult(True, t("html_ok"), content=html, engine="pandoc")
            except Exception as exc:  # noqa: BLE001
                fallback_msg = str(exc)
        else:
            fallback_msg = t("no_pandoc")

        try:
            html = _fallback_markdown_to_html(md_text)
            return ConvertResult(
                True,
                t("html_ok_detail", detail=fallback_msg),
                content=html,
                engine="fallback",
            )
        except Exception as exc:  # noqa: BLE001
            return ConvertResult(False, t("convert_fail_exc", exc=exc))


def copy_markdown_for_word(md_text: str) -> ConvertResult:
    """Convert Markdown and place native Word content (OMML formulas) on the clipboard."""
    from . import clipboard
    from .wordcom import copy_docx_to_clipboard, word_installed

    md_text = normalize_markdown(md_text)
    if not md_text.strip():
        return ConvertResult(False, t("empty"))

    with tempfile.TemporaryDirectory(prefix="md2word_") as tmp:
        docx_path = Path(tmp) / "clipboard.docx"
        docx_res = markdown_to_docx(md_text, docx_path)
        if not docx_res.success or not docx_path.is_file():
            return ConvertResult(False, docx_res.message or t("docx_fail"))

        if word_installed():
            try:
                copy_docx_to_clipboard(docx_path)
                return ConvertResult(
                    True,
                    t("copied_word", paste=clipboard.paste_shortcut()),
                    output_path=docx_path,
                    engine=f"{docx_res.engine}+word",
                )
            except Exception as exc:  # noqa: BLE001
                word_err = str(exc)
        else:
            word_err = t("no_word")

        html_res = markdown_to_html(md_text)
        if not html_res.success or not html_res.content:
            return ConvertResult(
                False,
                t("copy_fail", word=word_err, html=html_res.message),
            )
        plain = plain_text_from_markdown(md_text)
        formats = clipboard.set_rich_for_word(
            html=html_res.content,
            rtf=None,
            plain=plain,
        )
        return ConvertResult(
            True,
            t("copied_rich", formats=formats, word=word_err),
            content=html_res.content,
            engine=html_res.engine,
        )


def markdown_to_pdf(md_text: str, output_path: Path) -> ConvertResult:
    """Convert Markdown to PDF (Word export preferred, LibreOffice fallback)."""
    from .wordcom import export_docx_to_pdf, find_soffice, word_installed

    md_text = normalize_markdown(md_text)
    if not md_text.strip():
        return ConvertResult(False, t("empty"))

    output_path = Path(output_path)
    output_path.parent.mkdir(parents=True, exist_ok=True)

    with tempfile.TemporaryDirectory(prefix="md2word_") as tmp:
        docx_path = Path(tmp) / "export.docx"
        docx_res = markdown_to_docx(md_text, docx_path)
        if not docx_res.success or not docx_path.is_file():
            return ConvertResult(False, docx_res.message or t("docx_fail"))

        if word_installed():
            try:
                export_docx_to_pdf(docx_path, output_path)
                if output_path.is_file():
                    return ConvertResult(
                        True,
                        t("saved_path", path=output_path),
                        output_path=output_path,
                        engine="word",
                    )
            except Exception as exc:  # noqa: BLE001
                word_err = str(exc)
        else:
            word_err = t("no_word")

        soffice = find_soffice()
        if soffice:
            try:
                _libreoffice_to_pdf(soffice, docx_path, output_path)
                return ConvertResult(
                    True,
                    t("saved_lo", path=output_path),
                    output_path=output_path,
                    engine="libreoffice",
                )
            except Exception as exc:  # noqa: BLE001
                return ConvertResult(
                    False,
                    t("pdf_fail", word=word_err, lo=exc),
                )

        return ConvertResult(
            False,
            t("pdf_need", word=word_err),
        )


def _libreoffice_to_pdf(soffice: str, docx_path: Path, pdf_path: Path) -> None:
    outdir = pdf_path.parent
    cmd = [
        soffice,
        "--headless",
        "--norestore",
        "--convert-to",
        "pdf",
        "--outdir",
        str(outdir),
        str(docx_path),
    ]
    creationflags = 0
    if os.name == "nt":
        creationflags = getattr(subprocess, "CREATE_NO_WINDOW", 0)
    proc = subprocess.run(
        cmd,
        capture_output=True,
        text=True,
        encoding="utf-8",
        errors="replace",
        creationflags=creationflags,
        timeout=120,
    )
    if proc.returncode != 0:
        err = (proc.stderr or proc.stdout or "unknown error").strip()
        raise RuntimeError(err)
    produced = outdir / (docx_path.stem + ".pdf")
    if produced.resolve() != pdf_path.resolve():
        if pdf_path.exists():
            pdf_path.unlink()
        produced.replace(pdf_path)
    if not pdf_path.is_file():
        raise RuntimeError("LibreOffice 未生成 PDF 文件")


def markdown_to_rtf(md_text: str) -> ConvertResult:
    """Convert Markdown to RTF.

    Note: RTF is often a poor choice for math-heavy AI text because raw
    backslashes collide with RTF control words. Prefer HTML clipboard.
    """
    md_text = normalize_markdown(md_text)
    if not md_text.strip():
        return ConvertResult(False, t("empty"))

    pandoc = find_pandoc()
    if not pandoc:
        return ConvertResult(False, "RTF 转换需要安装 Pandoc。")

    with tempfile.TemporaryDirectory(prefix="md2word_") as tmp:
        tmp_dir = Path(tmp)
        md_path = tmp_dir / "input.md"
        rtf_path = tmp_dir / "out.rtf"
        md_path.write_text(md_text, encoding="utf-8")
        try:
            _run_pandoc(pandoc, md_path, rtf_path, "rtf")
            rtf = rtf_path.read_text(encoding="utf-8", errors="replace")
            return ConvertResult(True, "RTF 转换成功", content=rtf, engine="pandoc")
        except Exception as exc:  # noqa: BLE001
            return ConvertResult(False, f"RTF 转换失败：{exc}")


def markdown_to_plain(md_text: str) -> str:
    """Unicode-normalized plain text for clipboard fallback."""
    return plain_text_from_markdown(md_text)


def _inject_word_friendly_styles(html: str) -> str:
    """Ensure HTML has Word-friendly base styles."""
    style = """
<style>
  body {
    font-family: "Times New Roman", "SimSun", "宋体", serif;
    font-size: 12pt;
    line-height: 1.5;
    color: #222;
  }
  h1, h2, h3, h4, h5, h6 {
    font-family: "SimHei", "黑体", "Microsoft YaHei", sans-serif;
    font-weight: bold;
    color: #000;
    text-indent: 0;
    text-align: left;
    line-height: 1.5;
  }
  h1 { font-size: 18pt; margin: 14pt 0 8pt; }
  h2 { font-size: 16pt; margin: 12pt 0 8pt; }
  h3 { font-size: 14pt; margin: 10pt 0 6pt; }
  h4, h5, h6 { font-size: 12pt; margin: 10pt 0 6pt; }
  p {
    margin: 0;
    text-indent: 2em;
    text-align: justify;
    line-height: 1.5;
  }
  li p, td p, th p, blockquote p { text-indent: 0; }
  code, pre {
    font-family: Consolas, "Courier New", monospace;
    font-size: 10.5pt;
  }
  code {
    background: #f4f4f4;
    padding: 1px 4px;
  }
  pre {
    background: #f6f8fa;
    border: 1px solid #e1e4e8;
    padding: 10pt;
    white-space: pre-wrap;
    text-indent: 0;
  }
  blockquote {
    margin: 8pt 0;
    padding: 4pt 12pt;
    border-left: 3pt solid #ccc;
    color: #555;
    text-indent: 0;
  }
  table {
    border-collapse: collapse;
    margin: 8pt 0;
  }
  th, td {
    border: 1px solid #999;
    padding: 4pt 8pt;
    text-indent: 0;
  }
  th { background: #f0f0f0; }
  ul, ol { margin: 0 0 8pt 22pt; }
  li { margin: 2pt 0; text-indent: 0; }
  a { color: #0563c1; }
  hr { display: none; border: none; height: 0; margin: 0; }
  .math.display, math[display="block"] {
    display: block;
    text-align: center;
    text-indent: 0;
    margin: 8pt 0;
  }
  .math, math { font-family: "Cambria Math", "Times New Roman", serif; }
</style>
"""
    if re.search(r"</head>", html, flags=re.I):
        return re.sub(r"</head>", style + "</head>", html, count=1, flags=re.I)
    if re.search(r"<body", html, flags=re.I):
        return re.sub(
            r"<body([^>]*)>",
            r"<head><meta charset='utf-8'>" + style + r"</head><body\1>",
            html,
            count=1,
            flags=re.I,
        )
    return (
        "<html><head><meta charset='utf-8'>"
        + style
        + f"</head><body>{html}</body></html>"
    )


def _fallback_markdown_to_html(md_text: str) -> str:
    """Minimal Markdown → HTML without external deps."""
    lines = md_text.split("\n")
    out: list[str] = []
    in_code = False
    code_lang = ""
    in_ul = False
    in_ol = False
    para: list[str] = []

    def flush_para() -> None:
        nonlocal para
        if para:
            text = " ".join(para)
            out.append(f"<p>{_inline(text)}</p>")
            para = []

    def close_lists() -> None:
        nonlocal in_ul, in_ol
        if in_ul:
            out.append("</ul>")
            in_ul = False
        if in_ol:
            out.append("</ol>")
            in_ol = False

    for raw in lines:
        line = raw.rstrip("\n")
        fence = re.match(r"^```(\w*)\s*$", line)
        if fence:
            flush_para()
            close_lists()
            if not in_code:
                in_code = True
                code_lang = fence.group(1) or ""
                cls = f' class="language-{code_lang}"' if code_lang else ""
                out.append(f"<pre><code{cls}>")
            else:
                in_code = False
                out.append("</code></pre>")
            continue
        if in_code:
            out.append(
                line.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")
                + "\n"
            )
            continue

        if not line.strip():
            flush_para()
            close_lists()
            continue

        heading = re.match(r"^(#{1,6})\s+(.*)$", line)
        if heading:
            flush_para()
            close_lists()
            level = len(heading.group(1))
            out.append(f"<h{level}>{_inline(heading.group(2))}</h{level}>")
            continue

        if re.match(r"^---+\s*$", line) or re.match(r"^\*\*\*+\s*$", line):
            flush_para()
            close_lists()
            continue

        bq = re.match(r"^>\s?(.*)$", line)
        if bq:
            flush_para()
            close_lists()
            out.append(f"<blockquote><p>{_inline(bq.group(1))}</p></blockquote>")
            continue

        ul = re.match(r"^(\s*)[-*+]\s+(.*)$", line)
        if ul:
            flush_para()
            if in_ol:
                out.append("</ol>")
                in_ol = False
            if not in_ul:
                out.append("<ul>")
                in_ul = True
            out.append(f"<li>{_inline(ul.group(2))}</li>")
            continue

        ol = re.match(r"^(\s*)\d+\.\s+(.*)$", line)
        if ol:
            flush_para()
            if in_ul:
                out.append("</ul>")
                in_ul = False
            if not in_ol:
                out.append("<ol>")
                in_ol = True
            out.append(f"<li>{_inline(ol.group(2))}</li>")
            continue

        close_lists()
        para.append(line.strip())

    flush_para()
    close_lists()
    if in_code:
        out.append("</code></pre>")

    body = "\n".join(out)
    return _inject_word_friendly_styles(
        "<!DOCTYPE html><html><head><meta charset='utf-8'></head>"
        f"<body>{body}</body></html>"
    )


def _inline(text: str) -> str:
    text = text.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")
    text = re.sub(
        r"\$\$(.+?)\$\$",
        r'<span class="math display">\1</span>',
        text,
        flags=re.S,
    )
    text = re.sub(
        r"(?<!\$)\$(?!\$)(.+?)(?<!\$)\$(?!\$)",
        r'<span class="math inline">\1</span>',
        text,
        flags=re.S,
    )
    text = re.sub(r"`([^`]+)`", r"<code>\1</code>", text)
    text = re.sub(r"\*\*([^*]+)\*\*", r"<strong>\1</strong>", text)
    text = re.sub(r"__([^_]+)__", r"<strong>\1</strong>", text)
    text = re.sub(r"(?<!\*)\*([^*]+)\*(?!\*)", r"<em>\1</em>", text)
    text = re.sub(r"(?<!_)_([^_]+)_(?!_)", r"<em>\1</em>", text)
    text = re.sub(r"~~([^~]+)~~", r"<del>\1</del>", text)
    text = re.sub(r"\[([^\]]+)\]\(([^)]+)\)", r'<a href="\2">\1</a>', text)
    return text


def _fallback_markdown_to_docx(md_text: str, output_path: Path) -> None:
    """Basic DOCX export via python-docx."""
    from docx import Document
    from docx.shared import Pt

    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Times New Roman"
    style.font.size = Pt(12)

    lines = md_text.split("\n")
    in_code = False
    code_lines: list[str] = []

    def add_runs(paragraph, text: str) -> None:
        pattern = re.compile(r"(\*\*[^*]+\*\*|\*[^*]+\*|`[^`]+`|\[[^\]]+\]\([^)]+\)|\$\$[^$]+\$\$|\$[^$]+\$)")
        pos = 0
        for m in pattern.finditer(text):
            if m.start() > pos:
                paragraph.add_run(text[pos : m.start()])
            token = m.group(0)
            if token.startswith("**") and token.endswith("**"):
                run = paragraph.add_run(token[2:-2])
                run.bold = True
            elif token.startswith("*") and token.endswith("*"):
                run = paragraph.add_run(token[1:-1])
                run.italic = True
            elif token.startswith("`") and token.endswith("`"):
                run = paragraph.add_run(token[1:-1])
                run.font.name = "Consolas"
            elif token.startswith("$$") and token.endswith("$$"):
                run = paragraph.add_run(token[2:-2].strip())
                run.italic = True
                run.font.name = "Cambria Math"
            elif token.startswith("$") and token.endswith("$"):
                run = paragraph.add_run(token[1:-1])
                run.italic = True
                run.font.name = "Cambria Math"
            elif token.startswith("["):
                m2 = re.match(r"\[([^\]]+)\]\(([^)]+)\)", token)
                if m2:
                    paragraph.add_run(m2.group(1))
                else:
                    paragraph.add_run(token)
            else:
                paragraph.add_run(token)
            pos = m.end()
        if pos < len(text):
            paragraph.add_run(text[pos:])

    for raw in lines:
        line = raw.rstrip("\n")
        fence = re.match(r"^```", line)
        if fence:
            if not in_code:
                in_code = True
                code_lines = []
            else:
                in_code = False
                p = doc.add_paragraph()
                run = p.add_run("\n".join(code_lines))
                run.font.name = "Consolas"
                run.font.size = Pt(10)
            continue
        if in_code:
            code_lines.append(line)
            continue
        if not line.strip():
            continue
        heading = re.match(r"^(#{1,6})\s+(.*)$", line)
        if heading:
            level = min(len(heading.group(1)), 4)
            doc.add_heading(heading.group(2), level=level)
            continue
        if re.match(r"^---+\s*$", line) or re.match(r"^\*\*\*+\s*$", line):
            continue
        ul = re.match(r"^\s*[-*+]\s+(.*)$", line)
        if ul:
            p = doc.add_paragraph(style="List Bullet")
            add_runs(p, ul.group(1))
            continue
        ol = re.match(r"^\s*\d+\.\s+(.*)$", line)
        if ol:
            p = doc.add_paragraph(style="List Number")
            add_runs(p, ol.group(1))
            continue
        bq = re.match(r"^>\s?(.*)$", line)
        if bq:
            p = doc.add_paragraph()
            run = p.add_run(bq.group(1))
            run.italic = True
            continue
        p = doc.add_paragraph()
        add_runs(p, line.strip())

    doc.save(str(output_path))


def open_file(path: Path) -> None:
    """Open a file with the OS default app (Word for .docx)."""
    path = Path(path)
    if os.name == "nt":
        os.startfile(str(path))  # type: ignore[attr-defined]
    elif shutil.which("open"):
        subprocess.Popen(["open", str(path)])
    else:
        subprocess.Popen(["xdg-open", str(path)])
